from django.shortcuts import render,redirect,get_object_or_404
from django.http import HttpResponse, Http404, HttpResponseBadRequest
from .models import activo, UserData,Contador, AsignacionActivo,actaEntrega
from django.core.paginator import Paginator
from django.db.models import Q
from django.core.serializers import serialize
from django.core.mail import EmailMessage
from io import BytesIO
from datetime import datetime
from django.utils import timezone
import json  # Agregar esta línea
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate
import pandas as pd
import xml.etree.ElementTree as ET
import locale
from docx.oxml import parse_xml
from django.contrib.auth import login as django_login
from django.contrib.auth import login,logout,authenticate
from itertools import groupby
from operator import attrgetter 
from django.http import JsonResponse
from django.contrib.auth import authenticate, login 
from django.views.decorators.cache import never_cache
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth.models import User, Permission, Group
from django.contrib.auth import login, logout, authenticate
from django.db import IntegrityError
from django.db import models
from django.shortcuts import redirect
from django.contrib import messages

from django.core.exceptions import ValidationError

from datetime import datetime
from django.db.models import Count, Avg, ExpressionWrapper, DurationField
from django.core.mail import EmailMessage
from threading import Thread 
from django.urls import reverse  # Importar reverse
from .forms import ActaEntregaForm
from django.core.files.base import ContentFile



TEMPLATE_DIRS = (
    'os.path.join(BASE_DIR, "templates")'
)

#CREATE-DELETE-LIST AND UPDATE MODEL ACTIVE--->
#===============#===========================#

def index(request):
    # Verificar si el usuario es "totaladmin"
    is_totaladmin = request.user.is_authenticated and User.objects.filter(username="totaladmin").exists()
    print(f"¿Es totaladmin el usuario? {is_totaladmin}")
    
    # Activos más prestados
    activos_mas_prestados = (
        activo.objects.filter(asignacionactivo__fecha_devolucion__isnull=False)  # Activos con devoluciones
        .annotate(total_prestamos=Count("asignacionactivo"))
        .order_by("-total_prestamos")[:5]
    )

    # Usuarios más frecuentes
    usuarios_mas_frecuentes = (
        UserData.objects.annotate(total_prestamos=Count("asignacionactivo"))
        .order_by("-total_prestamos")[:5]
    )

    # Tiempo promedio de préstamo
    prestamos = AsignacionActivo.objects.filter(
        fecha_devolucion__isnull=False  # Solo préstamos devueltos
    ).annotate(
        duracion=ExpressionWrapper(
            models.F("fecha_devolucion") - models.F("fecha_asignacion"),
            output_field=DurationField()
        )
    )
    tiempo_promedio_prestamo = prestamos.aggregate(
        promedio=Avg("duracion")
    )["promedio"]

    # Convertir tiempo promedio a días si existe
    tiempo_promedio_prestamo_dias = None
    if tiempo_promedio_prestamo:
        tiempo_promedio_prestamo_dias = tiempo_promedio_prestamo.days

    context = {
        "is_totaladmin": is_totaladmin,
        "activos_mas_prestados": activos_mas_prestados,
        "usuarios_mas_frecuentes": usuarios_mas_frecuentes,
        "tiempo_promedio_prestamo": tiempo_promedio_prestamo_dias,
    }

    return render(request, "index.html", context)


def register(request):
    
    if request.method == 'GET':    
        return render(request, 'acces_user/register.html', {
            'form':UserCreationForm
        })
    
    else:
        
        if request.POST['password1'] == request.POST['password2']:
            try:
                user = User.objects.create_user(first_name=request.POST["first_name"],
                                                email=request.POST["email"],
                                                password=request.POST["password1"],
                                                username=request.POST["username"]
                                                )
                user.save()
                
                user = authenticate(
                    request,
                        first_name=request.POST["first_name"],
                        email=request.POST["email"],
                        password=request.POST["password1"],
                        username=request.POST["username"]
                )
                
                if user is not None:
                    login(request, user)  # Esto ya incluye el backend correcto
                    return redirect('index')

            except IntegrityError:
                
                return render(request, 'acces_user/register.html', {'form':UserCreationForm,"error": 'EL usuario ya existe'})

        return render(request, 'acces_user/register.html',{'form': UserCreationForm,"error": 'Contraseñas Incorrecta'})


def clouses(request):
    if request.method == 'POST':
        if 'logout' in request.POST:
            if request.user.is_authenticated:
                logout(request)
            return redirect(reverse('login'))  # Usar reverse para generar la URL del login
    return redirect('index')  # Si no es POST, redirigir al índice
    
    
@never_cache
def custom_login_view(request):
    if request.method == 'GET':
        return render(request, 'acces_user/login.html', {
            'form': AuthenticationForm
        })

    else:
        # Autenticar al usuario
        user = authenticate(
            request, username=request.POST['username'], password=request.POST['password']
        )

        if user is None:
            # Si la autenticación falla, mostrar error
            return render(request, 'acces_user/login.html', {
                'form': AuthenticationForm,
                'error': "La contraseña de la usuari@ es incorrecta"
            })
        else:
            # Si la autenticación es exitosa, iniciar sesión
            login(request, user)

            # Verificar si el usuario es "totaladmin"
            if user.username == "totaladmin":
                # Redirigir a una página especial si es "totaladmin"
                return redirect('/list_user')  # O la ruta que desees para "totaladmin"

            # Si no es "totaladmin", redirigir a /list_user
            return redirect('/list_user')





def data_user(request,iduser):

    usuario = get_object_or_404(UserData, pk=iduser)
    # trabajador = get_object_or_404(User, pk=user_id)
    
    context={
        'usuario':usuario,
        # 'trabajador': trabajador

    }
    
    return render(request,"data_genereitor/data_user.html",context)



def data_activo(request, idactivo):
    activo_obj = get_object_or_404(activo, pk=idactivo)
    asignacion = AsignacionActivo.objects.filter(activo=activo_obj, fecha_devolucion__isnull=True).first()
    usuario_asignado = asignacion.usuario if asignacion else None
    
    context = {
        'activo': activo_obj,
        'usuario_asignado': usuario_asignado,
    }
    
    return render(request, "data_genereitor/data_activo.html", context)


def detail_active(request, iduser,user_id):
    user = get_object_or_404(User, pk=user_id)
    
    
    usuario = get_object_or_404(UserData, pk=iduser)
    asignaciones = AsignacionActivo.objects.filter(usuario=usuario).order_by('fecha_asignacion')
    
    # Agrupar las asignaciones por fecha.
    asignaciones_agrupadas = {}
    for asignacion in asignaciones:
        fecha = asignacion.fecha_asignacion.strftime("%d/%m/%y")
        hora = asignacion.fecha_asignacion.strftime('%H:%M')
        fecha_devolucion = asignacion.fecha_devolucion.strftime("%d/%m/%Y %H:%M") if asignacion.fecha_devolucion else "-"
        clave = (fecha, hora)
        asignacion_datos = {'id': asignacion.id, 'activo': asignacion.activo, 'fecha_devolucion': fecha_devolucion}
        if clave in asignaciones_agrupadas:
            asignaciones_agrupadas[clave].append(asignacion_datos)
        else:
            asignaciones_agrupadas[clave] = [asignacion_datos]
            
    
    agrupaciones_lista = list(asignaciones_agrupadas.items())
    paginator = Paginator(agrupaciones_lista,6)


    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Construir el contexto para pasar a la plantilla.
    context = {
        'user':user,
        'usuario': usuario,
        'page_obj': page_obj
        
    }
        
    return render(request, "data_genereitor/detail_active.html", context)


def devolver_activo(request, asignacion_id):
    if request.method == 'POST':
        asignacion = get_object_or_404(AsignacionActivo, pk=asignacion_id)
        
        print(f"ID del activo: {asignacion}")

        # Verificar si el activo ya ha sido devuelto
        activo_devuelto = asignacion.fecha_devolucion is not None

        # Si el activo aún no ha sido devuelto, marcarlo como devuelto en la base de datos
        if not activo_devuelto:
            asignacion.fecha_devolucion = timezone.now()
            asignacion.save()  # Guardar el activo después de modificar su estado

                        # Obtener el activo asociado a esta asignación
            activo = asignacion.activo 
            
            # Modificar la declaración del activo a 'ACTIVO'
            activo.declaracion = 'ACTIVO'
            activo.save()  # Guardar el activo después de modificar su estado

        
        
        # Formatear la fecha y hora para enviarlas en la respuesta JSON
        fecha_devolucion = asignacion.fecha_devolucion.strftime("%d/%m/%Y")
        hora_devolucion = asignacion.fecha_devolucion.strftime("%H:%M")
        
        # Serializar solo los campos necesarios de la asignación
        asignacion_serializada = serialize('json', [asignacion], fields=('pk', 'fecha_devolucion'))

        # Cargar la cadena JSON serializada en un objeto Python
        asignacion_serializada_json = json.loads(asignacion_serializada)

        # Acceder a los campos serializados
        asignacion_campos = asignacion_serializada_json[0]['fields']
        
        return JsonResponse({'fecha_devolucion': fecha_devolucion,
                             'hora_devolucion': hora_devolucion,
                             'asignacion': asignacion_campos,
                             'activo_devuelto': not activo_devuelto
                              })
                
    else:
        # return redirect('detail_active', iduser=asignacion.id, fecha_devolucion=asignacion.fecha_devolucion)
        return JsonResponse({'error': 'Método no permitido'}, status=405)




def acta_entrega (request):
    return render(request, "acta_entrega.html")


def list_actives (request):
    if request.method=='POST':
        word = request.POST.get('keyword')
        list = activo.objects.all()

        if word is not None:
            resultado_busqueda = list.filter(
                 
                Q(id__icontains=word) |
                Q(serie__icontains=word)|
                Q(codigo_inventario__icontains=word)|
                Q(categoria__icontains=word) 
                
            )


            paginator = Paginator(resultado_busqueda, 8)  # 10 usuarios por página
        else:
            paginator = Paginator(list, 8)
    else:
        activos = activo.objects.all()
        paginator = Paginator(activos, 8)  # 10 usuarios por página

    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    datos = {'page_obj': page_obj}

    return render(request, "crud_actives/list_actives.html",datos)



def add_actives(request):
    if request.method == 'POST':
        # Captura de los campos del formulario
        data = {
            'sede': request.POST.get('sede', ''),
            'pabellon': request.POST.get('pabellon', ''),
            'tipo_ambiente': request.POST.get('tipo_ambiente', ''),
            'ambiente': request.POST.get('ambiente', ''),
            'direccion': request.POST.get('direccion', ''),
            'distrito': request.POST.get('distrito', ''),
            'serie': request.POST.get('serie', ''),
            'codigo_inventario': request.POST.get('codigo_inventario', ''),
            'categoria': request.POST.get('categoria', ''),
            'marca': request.POST.get('marca', ''),
            'descripcion': request.POST.get('descripcion', ''),
            'modelo': request.POST.get('modelo', ''),
            'hostname': request.POST.get('hostname', ''),
            'estado': request.POST.get('estado', ''),
            'renta': request.POST.get('renta', ''),
            'contrato': request.POST.get('contrato', ''),
            'estado_renta': request.POST.get('estado_renta', ''),
            'proveedor': request.POST.get('proveedor', ''),
            'devolucion': request.POST.get('devolucion', ''),
            'declaracion': request.POST.get('declaracion', 'ACTIVO'),
        }

        errors = {}

        try:
            # Validar campos requeridos
            required_fields = ['sede', 'categoria', 'estado', 'codigo_inventario']
            for field in required_fields:
                if not data[field]:
                    errors[field] = 'Este campo es obligatorio.'

            # Validar valores únicos
            if activo.objects.filter(codigo_inventario=data['codigo_inventario']).exists():
                errors['codigo_inventario'] = 'El código de inventario ya existe.'

            # Si hay errores, devolver formulario con datos y mensajes de error
            if errors:
                return render(request, "crud_actives/add_actives.html", {
                    'errors': errors,
                    'activo': data
                })

            # Guardar activo
            active = activo(**data)
            active.save()

            return redirect('list_actives')
        except Exception as e:
            return render(request, "crud_actives/add_actives.html", {
                'error': f'Error al guardar el activo: {e}',
                'activo': data
            })

    return render(request, "crud_actives/add_actives.html")




def update_actives(request, idactive):
    try:
        # Obtener el registro específico
        activo_instance = activo.objects.get(id=idactive)

        if request.method == 'POST':
            # Actualizar todos los campos del modelo desde los datos del formulario
            activo_instance.sede = request.POST.get('sede', activo_instance.sede)
            activo_instance.pabellon = request.POST.get('pabellon', activo_instance.pabellon)
            activo_instance.tipo_ambiente = request.POST.get('tipo_ambiente', activo_instance.tipo_ambiente)
            activo_instance.ambiente = request.POST.get('ambiente', activo_instance.ambiente)
            activo_instance.direccion = request.POST.get('direccion', activo_instance.direccion)
            activo_instance.distrito = request.POST.get('distrito', activo_instance.distrito)
            activo_instance.serie = request.POST.get('serie', activo_instance.serie)
            activo_instance.codigo_inventario = request.POST.get('codigo_inventario', activo_instance.codigo_inventario)
            activo_instance.categoria = request.POST.get('categoria', activo_instance.categoria)
            activo_instance.marca = request.POST.get('marca', activo_instance.marca)
            activo_instance.descripcion = request.POST.get('descripcion', activo_instance.descripcion)
            activo_instance.modelo = request.POST.get('modelo', activo_instance.modelo)
            activo_instance.hostname = request.POST.get('hostname', activo_instance.hostname)
            activo_instance.estado = request.POST.get('estado', activo_instance.estado)
            activo_instance.renta = request.POST.get('renta', activo_instance.renta)
            activo_instance.contrato = request.POST.get('contrato', activo_instance.contrato)
            activo_instance.estado_renta = request.POST.get('estado_renta', activo_instance.estado_renta)
            activo_instance.proveedor = request.POST.get('proveedor', activo_instance.proveedor)
            activo_instance.devolucion = request.POST.get('devolucion', activo_instance.devolucion)
            activo_instance.declaracion = request.POST.get('declaracion', activo_instance.declaracion)

            # Guardar los cambios
            activo_instance.save()
            return redirect('list_actives')  # Redirigir a la lista de activos

        else:
            # Enviar el registro actual a la plantilla para editar
            datos = {'activo_instance': activo_instance}
            return render(request, "crud_actives/update_actives.html", datos)

    except activo.DoesNotExist:
        # Si el registro no existe, manejar el error y mostrar un mensaje
        return render(request, "crud_actives/update_actives.html", {
            'error': 'El activo no existe.'
        })



def delete_actives(request, idactive):
    try:
        # Intenta obtener el activo de la base de datos usando el nombre correcto
        activo_obj = get_object_or_404(activo, id=idactive)

        if request.method == 'POST':
            # Confirmación de eliminación
            activo_obj.delete()
            
            # Agregar mensaje de éxito a la sesión
            request.session['alert'] = 'Activo eliminado exitosamente.'

            # Redirigir a la lista de activos
            return redirect('list_actives')

        else:
            # Si no es un POST, pasa el activo a la plantilla para la confirmación
            return render(request, "crud_actives/delete_actives.html", {'activo': activo_obj})

    except activo.DoesNotExist:
        # Si el activo no existe, pasamos un mensaje de error a la plantilla
        return render(request, "crud_actives/delete_actives.html", {'error': 'Activo no encontrado'})

    except Exception as e:
        # Manejo de otros errores
        return render(request, "crud_actives/delete_actives.html", {'error': str(e)})


#**************************#*************************************************#

    #LIST - CREATE - UPDATE AND DELETE MODEL USER

#************************#***************************************************#
def list_user(request):
    if request.method == 'POST':
        palabra = request.POST.get('keyword')
        lista = UserData.objects.all()

        if palabra is not None:
            resultado_busqueda = lista.filter(
                Q(id__icontains=palabra) |
                Q(area_de_trabajo__icontains=palabra) |
                Q(sub_area_de_trabajo__icontains=palabra) |
                Q(codigo_de_personal__icontains=palabra) |
                Q(apellidos_y_nombres_adryan__icontains=palabra) |
                Q(tipo_de_jornada_adryan__icontains=palabra)   |
                Q(correo_corporativo_adryan__icontains=palabra) 


            )
            
            paginator = Paginator(resultado_busqueda, 8)  # 10 usuarios por página
        else:
            paginator = Paginator(lista, 8)
    else:
        usuarios = UserData.objects.all()
        paginator = Paginator(usuarios, 8)  # 10 usuarios por página


    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    datos = {'page_obj': page_obj}

    return render(request, "crud_users/list_user.html", datos)


def add_user(request):
    if request.method == 'POST':
        # Captura de los campos del formulario
        data = {
            
            'planilla': request.POST.get('planilla', ''),
            'unidad_de_negocio': request.POST.get('unidad_de_negocio', ''),
            'area_de_trabajo': request.POST.get('area_de_trabajo', ''),
            'sub_area_de_trabajo': request.POST.get('sub_area_de_trabajo', ''),
            'ubicacion_fisica': request.POST.get('ubicacion_fisica', ''),
            'local': request.POST.get('local', ''),
            'naturaleza_de_puesto': request.POST.get('naturaleza_de_puesto', ''),
            'nomenclatura_de_puesto': request.POST.get('nomenclatura_de_puesto', ''),
            'tipo_de_puesto': request.POST.get('tipo_de_puesto', ''),
            'motivo_de_alta': request.POST.get('motivo_de_alta', ''),
            'codigo_de_personal': request.POST.get('codigo_de_personal', 0),
            'apellidos_y_nombres_adryan': request.POST.get('apellidos_y_nombres_adryan', ''),
            'genero_adryan': request.POST.get('genero_adryan', ''),
            'fecha_de_ingreso': request.POST.get('fecha_de_ingreso', '1900-01-01'),
            'tipo_de_contrato_adryan': request.POST.get('tipo_de_contrato_adryan', ''),
            'tipo_de_jornada_adryan': request.POST.get('tipo_de_jornada_adryan', ''),
            'nacionalidad_adryan': request.POST.get('nacionalidad_adryan', ''),
            'tiempo_de_servicio_anios': request.POST.get('tiempo_de_servicio_anios', 0),
            'tiempo_de_servicio_meses': request.POST.get('tiempo_de_servicio_meses', 0),
            'tiempo_de_servicio_dias': request.POST.get('tiempo_de_servicio_dias', 0),
            'correo_corporativo_adryan': request.POST.get('correo_corporativo_adryan', ''),
            'fecha_de_nacimiento_adryan': request.POST.get('fecha_de_nacimiento_adryan', '1900-01-01'),
            'edad_adryan': request.POST.get('edad_adryan', 0),
            'generacion_adryan': request.POST.get('generacion_adryan', ''),
            'jefe_inmediato_jerarquico': request.POST.get('jefe_inmediato_jerarquico', ''),
            'reemplaza_a': request.POST.get('reemplaza_a', ''),
        }

        # Validaciones
        try:
            # Validar que los campos requeridos estén completos
            required_fields = ['planilla', 'unidad_de_negocio', 'area_de_trabajo', 'codigo_de_personal', 'apellidos_y_nombres_adryan']
            if not all(data[field] for field in required_fields):
                raise ValidationError('Por favor complete todos los campos obligatorios.')

            # Validar fecha de ingreso y fecha de nacimiento
            try:
                datetime.strptime(data['fecha_de_ingreso'], '%Y-%m-%d')
                datetime.strptime(data['fecha_de_nacimiento_adryan'], '%Y-%m-%d')
            except ValueError:
                raise ValidationError('Las fechas deben tener el formato AAAA-MM-DD.')

            # Validar valores numéricos
            if int(data['tiempo_de_servicio_anios']) > 100 or int(data['tiempo_de_servicio_meses']) > 12 or int(data['tiempo_de_servicio_dias']) > 31:
                raise ValidationError('Valores de tiempo de servicio inválidos.')

            if int(data['edad_adryan']) > 100:
                raise ValidationError('La edad no puede superar los 100 años.')

            # Guardar usuario
            user = UserData(**data)  # Asegúrate de que User sea el nombre de tu modelo.
            user.save()
            return redirect('list_user')
        
        except ValidationError as e:
            return render(request, "crud_users/add_user.html", {'error': str(e), 'usuario': user})
        except Exception as e:
            return render(request, "crud_users/add_user.html", {'error': f'Error al guardar el usuario: {e}', 'usuario': user})

    return render(request, "crud_users/add_user.html")



def update_user(request, iduser):
    if request.method == 'POST':
        # Obtener el usuario correspondiente
        user = get_object_or_404(UserData, id=iduser)
        
        # Actualizar los campos con los datos del formulario
        user.planilla = request.POST.get('planilla', user.planilla)
        user.unidad_de_negocio = request.POST.get('unidad_de_negocio', user.unidad_de_negocio)
        user.area_de_trabajo = request.POST.get('area_de_trabajo', user.area_de_trabajo)
        user.sub_area_de_trabajo = request.POST.get('sub_area_de_trabajo', user.sub_area_de_trabajo)
        user.ubicacion_fisica = request.POST.get('ubicacion_fisica', user.ubicacion_fisica)
        user.local = request.POST.get('local', user.local)
        user.naturaleza_de_puesto = request.POST.get('naturaleza_de_puesto', user.naturaleza_de_puesto)
        user.nomenclatura_de_puesto = request.POST.get('nomenclatura_de_puesto', user.nomenclatura_de_puesto)
        user.tipo_de_puesto = request.POST.get('tipo_de_puesto', user.tipo_de_puesto)
        user.motivo_de_alta = request.POST.get('motivo_de_alta', user.motivo_de_alta)
        user.codigo_de_personal = request.POST.get('codigo_de_personal', user.codigo_de_personal)
        user.apellidos_y_nombres_adryan = request.POST.get('apellidos_y_nombres_adryan', user.apellidos_y_nombres_adryan)
        user.genero_adryan = request.POST.get('genero_adryan', user.genero_adryan)
        user.tipo_de_contrato_adryan = request.POST.get('tipo_de_contrato_adryan', user.tipo_de_contrato_adryan)
        user.tipo_de_jornada_adryan = request.POST.get('tipo_de_jornada_adryan', user.tipo_de_jornada_adryan)
        user.nacionalidad_adryan = request.POST.get('nacionalidad_adryan', user.nacionalidad_adryan)
        user.tiempo_de_servicio_anios = request.POST.get('tiempo_de_servicio_anios', user.tiempo_de_servicio_anios)
        user.tiempo_de_servicio_meses = request.POST.get('tiempo_de_servicio_meses', user.tiempo_de_servicio_meses)
        user.tiempo_de_servicio_dias = request.POST.get('tiempo_de_servicio_dias', user.tiempo_de_servicio_dias)
        user.correo_corporativo_adryan = request.POST.get('correo_corporativo_adryan', user.correo_corporativo_adryan)
        user.edad_adryan = request.POST.get('edad_adryan', user.edad_adryan)
        user.generacion_adryan = request.POST.get('generacion_adryan', user.generacion_adryan)
        user.jefe_inmediato_jerarquico = request.POST.get('jefe_inmediato_jerarquico', user.jefe_inmediato_jerarquico)
        user.reemplaza_a = request.POST.get('reemplaza_a', user.reemplaza_a)
        
        user.save()  # Guardar los cambios
        
        # Mostrar un mensaje de éxito
        messages.success(request, '¡Usuario actualizado correctamente!')
        return redirect('list_user')
    
    else:
        # Mostrar el formulario con datos del usuario seleccionado
        usuarios = UserData.objects.all()
        user = get_object_or_404(UserData, id=iduser)
        return render(request, "crud_users/update_user.html", {'usuarios': usuarios, 'usuario': user})


def delete_user(request, iduser):
    try:
        # Obtiene el usuario a eliminar
        usuario = get_object_or_404(UserData, id=iduser)
        
        if request.method == 'POST':
            # Confirmación de eliminación
            usuario.delete()
            # Redirige y muestra una alerta de éxito
            request.session['alert'] = 'Usuario eliminado exitosamente.'
            return redirect('list_user')

        else:
            # Si no es POST, muestra la confirmación
            datos = {'usuario': usuario}
            return render(request, "crud_users/delete_user.html", datos)
    
    except Exception as e:
        # Manejo de errores
        datos = {'error': str(e)}
        return render(request, "crud_users/delete_user.html", datos)

###############REPORTES DE USUARIOS###################

#============FUNCIONES ANIDADADES DEL GENERADOR DEL ACTA DE ENTREGA====#


def generar_acta_entrega(request, iduser, user_id):
    try:
        user = obtener_trabajador(user_id)
        usuario = obtener_usuario(iduser)
        contador = obtener_contador()
        keyword = request.POST.get('keyword', '')
        
        # Filtrar activos disponibles
        activos_list = obtener_activos().exclude(declaracion='ASIGNADO')
        if keyword:
            activos_list = activos_list.filter(
                Q(categoria__icontains=keyword) |
                Q(descripcion__icontains=keyword) |
                Q(modelo__icontains=keyword) |
                Q(estado__icontains=keyword)
            )

        # Paginación
        paginator = Paginator(activos_list, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        # Manejo de POST para generar acta
        if request.method == 'POST' and 'generar_acta' in request.POST:
            activos_seleccionados = request.POST.getlist('activo')
            if not activos_seleccionados:
                messages.error(request, "Debe seleccionar al menos un activo para generar el acta.")
                return redirect('acta_entrega')  # Ajusta el nombre según tu URL

            # Procesar activos seleccionados
            activos_entregados = activo.objects.filter(pk__in=activos_seleccionados)
            asignaciones_agrupadas = agrupar_asignaciones_por_fecha(
                AsignacionActivo.objects.filter(usuario=usuario)
            )

            for activo_entregado in activos_entregados:
                asignar_activo_a_usuario(usuario, activo_entregado)
                activo_entregado.declaracion = 'ASIGNADO'
                activo_entregado.save()

            # Crear documento
            doc = crear_documento(usuario, contador, asignaciones_agrupadas, activos_entregados, user)
            acta_bytes = document_to_bytes(doc)  # Convierte el documento a bytes
            
            # Guardar el archivo en la base de datos
            acta = actaEntrega(usuario=usuario, archivo_word=ContentFile(acta_bytes, name=f'acta_entrega_{contador}.docx'))
            acta.save()
        
        
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=acta_entrega_{contador}.docx'
            response.write(document_to_bytes(doc))
            actualizar_contador()
            
            # Obtener correos electrónicos
            correo_corporativo_adryan = usuario.correo_corporativo_adryan  # Desde UserData
            correo_trabajador = user.email  # Desde user
            destinatarios = [correo_corporativo_adryan, correo_trabajador]
            # Extraer los valores necesarios
            apellidos_y_nombres_adryan = usuario.apellidos_y_nombres_adryan
            first_name = user.first_name
            # Enviar correos en segundo plano
            Thread(
                target=enviar_correos_en_segundo_plano, 
                args=(destinatarios, acta_bytes, contador, apellidos_y_nombres_adryan, first_name)
            ).start()

            return response

        # Renderizar la página con el formulario
        context = {
            'user': user,
            'usuario': usuario,
            'page_obj': page_obj,
            'contador': contador,
            'keyword': keyword,
        }
        return render(request, 'actas_entrega/acta_entrega.html', context)

    except Exception as e:
        messages.error(request, f"Ha ocurrido un error: {str(e)}")
        return redirect('acta_entrega')  # Ajusta según tu configuración


def enviar_correos_en_segundo_plano(destinatarios, acta_bytes, contador, apellidos_y_nombres_adryan, first_name):
    """
    Envía los correos con el acta adjunta en un hilo separado.
    """
    try:
        email = EmailMessage(
            subject=f"Acta de entrega #{contador}",
            body=f"""\
                Estimado(a) {apellidos_y_nombres_adryan},

                Adjunto encontrará el acta de entrega generada automáticamente. 
                Para garantizar una mejor visualización del contenido, le recomendamos descargar el archivo adjunto y abrirlo en un visor compatible con documentos Word (.docx).

                Si tiene alguna consulta o necesita asistencia, no dude en ponerse en contacto.

                Saludos cordiales,
                {first_name}
                            """,
            from_email='quitoespirituleonardo@gmail.com',
            to=destinatarios,
        )

        # Adjuntar archivo
        email.attach(
            f"acta_entrega_{contador}.docx",
            acta_bytes,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        email.send()
    except Exception as e:
        # Aquí puedes registrar el error si falla el envío del correo
        print(f"Error enviando correo: {e}")




def obtener_usuario(iduser):
    return get_object_or_404(UserData, pk=iduser)
    

def obtener_trabajador(user_id):
    return get_object_or_404(User,pk=user_id)


    
def obtener_activos():
    return activo.objects.all()
    
    
    
def obtener_contador():
    contador_obj, created = Contador.objects.get_or_create(id=1)
    return contador_obj.valor



def agrupar_asignaciones_por_fecha(asignaciones):
    asignaciones_agrupadas = []
    for fecha, asignaciones_fecha in groupby(asignaciones, key=attrgetter('fecha_asignacion.date')):
        asignaciones_agrupadas.append((fecha, list(asignaciones_fecha)))
    return asignaciones_agrupadas
    
    
#VARIABLES DE ASIGNACION
def asignar_activo_a_usuario(usuario, activo_entregado):
    asignacion = AsignacionActivo(usuario=usuario, activo=activo_entregado)
   
    activo_entregado.declaracion = 'ASIGNADO'
    asignacion.save()
    
    
def actualizar_contador():
    contador_obj = Contador.objects.get(id=1)
    contador_obj.valor += 1
    contador_obj.save()



def convertir_campos_a_mayusculas(activos_entregados):
    for activo_entregado in activos_entregados:
        for field in ['descripcion', 'marca', 'serie', 'estado', 'estado']:
            setattr(activo_entregado, field, getattr(activo_entregado, field).upper())
        activo_entregado.save()



def crear_documento(usuario, contador, asignaciones_agrupadas, activos_entregados,trabajador):
    
    template_path = './INVENTARIO/templates/Documents/docxtemplate.docx'
    doc = DocxTemplate(template_path)
        # Renderizar la plantilla con los datos
    context = {
            'first_name': trabajador.first_name.upper(),
            'apellidos_y_nombres_adryan': usuario.apellidos_y_nombres_adryan.upper(),
            'nomenclatura_de_puesto': usuario.nomenclatura_de_puesto.upper(),
            'area_de_trabajo': usuario.area_de_trabajo.upper(),
            'contador': contador,
            'asignaciones_agrupadas': asignaciones_agrupadas,
        }
    doc.render(context)
    doc.add_paragraph("A continuación, se detallan los recursos informáticos")

        # Añadir una tabla para los activos entregados
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = True

    for cell in table.columns:
            cell.width = 5000

        # Encabezados de tabla
    headings = table.rows[0].cells
    headings[0].text = 'CANTIDAD'
    headings[1].text = 'DESCRIPCION'
    headings[2].text = 'MODELO'
    headings[3].text = 'CODIGO PATRIMONIO'
    headings[4].text = 'SERIE'
    headings[5].text = 'OBSERVACIONES'

        # Cambiar el color del fondo del encabezado a negro
    for cell in headings:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Color blanco
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w'))))  # Color de fondo en negro

    for cell in headings:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear al centro

    for activo_entregado in activos_entregados:
            row_cells = table.add_row().cells
            for i in range(6):
                row_cells[i].text = "1" if i == 0 else getattr(activo_entregado, ['descripcion', 'marca', 'serie', 'estado', 'estado'][i-1])
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("El presente documento deja constancia de que los recursos fueron entregados y recepcionados sin observaciones.")
    doc.add_paragraph("IMPORTANTE:")
    doc.add_paragraph("Por favor revisen la politica de activos fijos, pues este equipo se encuentra bajo su responsabilidad. https://conecta.continental.edu.pe/")

        # Obtener y formatear la fecha actual
    locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')
    fecha_actual = datetime.now()
    fecha_formateada = fecha_actual.strftime('Huancayo,' + '%A, %d de %B de %Y')
    texto_concatenado = fecha_formateada

        # Agregar el texto concatenado al documento
    parrafo = doc.add_paragraph(texto_concatenado)
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    return doc



def document_to_bytes(doc):
    document_buffer = BytesIO()
    doc.save(document_buffer)
    document_buffer.seek(0)
    return document_buffer.read()


###########    aqui termina         ###############


#----------------------EXPORTAR E IMPORTAR ARCHIVOS--------------------#
def import_usuarios(request):
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_usuarios')

        if excel_file and excel_file.name.endswith('.xlsx'):
            # Leer el archivo Excel
            df = pd.read_excel(excel_file)

            if all(col in df.columns for col in ['Planilla', 'Unidad de negocio', 'Área de trabajo', 'Sub Área de trabajo',
                                                 'Ubicación Física', 'Local ', 'Naturaleza de puesto', 'Nomenclatura de puesto',
                                                 'Tipo de puesto', 'Motivo de alta', 'Código de personal', 
                                                 'Apellidos y nombres (ADRYAN)', 'Género (ADRYAN)', 'Fecha de ingreso', 
                                                 'Tipo de contrato (ADRYAN)', 'Tipo de jornada (ADRYAN)', 'Nacionalidad (ADRYAN)', 
                                                 'Tiempo de servicio (años)', 'Tiempo de servicio (meses)', 'Tiempo de servicio (días)', 
                                                 'Correo corporativo (ADRYAN)', 'Fecha de nacimiento (ADRYAN)', 'Edad (ADRYAN)', 
                                                 'Generación (ADRYAN)', 'Jefe inmediato jerárquico', 'Reemplaza a']):
                
                duplicados = []  # Lista para duplicados
                importados = []  # Lista para importados correctamente

                # Brindar formato deseado a las fechas
                df['Fecha de ingreso'] = pd.to_datetime(df['Fecha de ingreso'], format='%d/%m/%Y', errors='coerce').dt.strftime('%Y-%m-%d')
                df['Fecha de nacimiento (ADRYAN)'] = pd.to_datetime(df['Fecha de nacimiento (ADRYAN)'], format='%d/%m/%Y', errors='coerce').dt.strftime('%Y-%m-%d')

                for index, row in df.iterrows():
                    # Validar duplicados
                    correo = row['Correo corporativo (ADRYAN)']
                    nombre = row['Apellidos y nombres (ADRYAN)']

                    if UserData.objects.filter(apellidos_y_nombres_adryan=nombre, correo_corporativo_adryan=correo).exists():
                        duplicados.append({
                            'fila': index + 2,  # +2 por encabezado y base 0 de Python
                            'nombre': nombre,
                            'correo': correo
                        })
                        continue  # Saltar registro duplicado

                    # Crear nuevo usuario si no es duplicado
                    usuario = UserData(
                        planilla=row['Planilla'],
                        unidad_de_negocio=row['Unidad de negocio'],
                        area_de_trabajo=row['Área de trabajo'],
                        sub_area_de_trabajo=row['Sub Área de trabajo'],
                        ubicacion_fisica=row['Ubicación Física'],
                        local=row['Local '],
                        naturaleza_de_puesto=row['Naturaleza de puesto'],
                        nomenclatura_de_puesto=row['Nomenclatura de puesto'],
                        tipo_de_puesto=row['Tipo de puesto'],
                        motivo_de_alta=row['Motivo de alta'],
                        codigo_de_personal=row['Código de personal'],
                        apellidos_y_nombres_adryan=nombre,
                        genero_adryan=row['Género (ADRYAN)'],
                        fecha_de_ingreso=row['Fecha de ingreso'],
                        tipo_de_contrato_adryan=row['Tipo de contrato (ADRYAN)'],
                        tipo_de_jornada_adryan=row['Tipo de jornada (ADRYAN)'],
                        nacionalidad_adryan=row['Nacionalidad (ADRYAN)'],
                        tiempo_de_servicio_anios=row['Tiempo de servicio (años)'],
                        tiempo_de_servicio_meses=row['Tiempo de servicio (meses)'],
                        tiempo_de_servicio_dias=row['Tiempo de servicio (días)'],
                        correo_corporativo_adryan=correo,
                        fecha_de_nacimiento_adryan=row['Fecha de nacimiento (ADRYAN)'],
                        edad_adryan=row['Edad (ADRYAN)'],
                        generacion_adryan=row['Generación (ADRYAN)'],
                        jefe_inmediato_jerarquico=row['Jefe inmediato jerárquico'],
                        reemplaza_a=row['Reemplaza a']
                    )
                    usuario.save()
                    importados.append({
                        'fila': index + 2,
                        'nombre': nombre,
                        'correo': correo
                    })

                # Mostrar mensajes
                if importados:
                    importados_msg = "\n".join(
                        [f"Fila {imp['fila']}: {imp['nombre']} ({imp['correo']})" for imp in importados]
                    )
                    messages.success(request, f"Usuarios importados correctamente:\n{importados_msg}")

                if duplicados:
                    duplicados_msg = "\n".join(
                        [f"Duplicado en fila {dup['fila']}: {dup['nombre']} ({dup['correo']})" for dup in duplicados]
                    )
                    messages.warning(request, f"Se encontraron duplicados:\n{duplicados_msg}")

                return redirect('list_user')

            else:
                messages.error(request, "El archivo Excel no contiene las columnas necesarias.")
        else:
            messages.error(request, "Por favor, suba un archivo Excel con extensión .xlsx.")

    return render(request, 'reports/Importar/usuarios.html')





def import_activos(request):
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_activos')

        if excel_file and excel_file.name.endswith('.xlsx'):
            # Leer el archivo Excel
            df = pd.read_excel(excel_file)

            if all(col in df.columns for col in ['SEDE', 'PABELLÓN', 'TIPO DE AMBIENTE', 'AMBIENTE', 'DIRECCIÓN', 
                                                 'DISTRITO', 'SERIE', 'CÓDIGO DE INVENTARIO', 'CATEGORIA', 
                                                 'MARCA', 'DESCRIPCIÓN', 'MODELO', 'HOSTNAME', 'ESTADO', 
                                                 'RENTA', 'CONTRATO', 'ESTADO DE RENTA', 'PROVEEDOR', 'DEVOLUCIÓN']):
                
                duplicados = []  # Lista para almacenar los duplicados
                importados = []  # Lista para almacenar los registros importados correctamente

                for index, row in df.iterrows():
                    # Validar duplicados
                    serie = row['SERIE']
                    codigo_inventario = row['CÓDIGO DE INVENTARIO']

                    if activo.objects.filter(serie=serie, codigo_inventario=codigo_inventario).exists():
                        duplicados.append({
                            'fila': index + 2,  # +2 para mostrar la fila real en el Excel (por el encabezado)
                            'serie': serie,
                            'codigo_inventario': codigo_inventario
                        })
                        continue  # Saltar este registro si ya existe

                    # Crear nuevo activo si no es duplicado
                    active = activo(
                        sede=row['SEDE'],
                        pabellon=row['PABELLÓN'],
                        tipo_ambiente=row['TIPO DE AMBIENTE'],
                        ambiente=row['AMBIENTE'],
                        direccion=row['DIRECCIÓN'],
                        distrito=row['DISTRITO'],
                        serie=serie,
                        codigo_inventario=codigo_inventario,
                        categoria=row['CATEGORIA'],
                        marca=row['MARCA'],
                        descripcion=row['DESCRIPCIÓN'],
                        modelo=row['MODELO'],
                        hostname=row['HOSTNAME'],
                        estado=row['ESTADO'],
                        renta=row['RENTA'],
                        contrato=row['CONTRATO'],
                        estado_renta=row['ESTADO DE RENTA'],
                        proveedor=row['PROVEEDOR'],
                        devolucion=row['DEVOLUCIÓN']
                    )
                    active.save()
                    importados.append({
                        'fila': index + 2,
                        'serie': serie,
                        'codigo_inventario': codigo_inventario
                    })

                # Mostrar mensajes
                if importados:
                    importados_msg = "\n".join(
                        [f"Fila {imp['fila']}: Serie {imp['serie']}, Código {imp['codigo_inventario']}" for imp in importados]
                    )
                    messages.success(request, f"Activos importados correctamente:\n{importados_msg}")

                if duplicados:
                    duplicados_msg = "\n".join(
                        [f"Duplicado en fila {dup['fila']}: Serie {dup['serie']}, Código {dup['codigo_inventario']}" for dup in duplicados]
                    )
                    messages.warning(request, f"Se encontraron duplicados:\n{duplicados_msg}")

                return redirect('activos')

            else:
                messages.error(request, "El archivo Excel no tiene las columnas necesarias.")
        else:
            messages.error(request, "Por favor, suba un archivo Excel con extensión .xlsx.")

    return render(request, 'reports/Importar/activos.html')



def export(request):
   return render(request, "reports/export.html")




def ver_actas_entrega(request, user_id):
    try:
        # Obtener todas las actas de entrega
        usuario = get_object_or_404(UserData, pk=user_id)
        
        # Filtrar las actas que pertenecen a este usuario
        actas = actaEntrega.objects.filter(usuario=usuario)

        context = {
            'actas': actas,
            'usuario': usuario,  # Pasar el usuario al contexto
        }
        return render(request, 'Informacion_Actas/ver_actas.html', context)
    except Exception as e:
        messages.error(request, f"Error al obtener las actas de entrega: {str(e)}")
        return redirect('home')  # Ajusta la URL de destino si es necesario



def delete_ActaEntrega(request, idacta):
    if request.method == 'POST':
        try:
            actas_entrega = acta_entrega.objects.get(pk=idacta)
            # Elimina el documento de la base de datos y del sistema de archivos
            actas_entrega.delete()
        except acta_entrega.DoesNotExist:
            # Maneja el caso en el que el documento no existe
            raise Http404("El acta de entrega no existe.")
    return redirect('info_acta', iduser=actas_entrega.Users.id)
    

#------------------Envio de correo Electronico---------#

def envio_email(request):

    if request.method == 'POST':
        email = request.POST.get('email', '')

        if email:
            subject = 'Firma Acta Entrega'
            message = 'Estimado colaborador reciba un cordial saludo de parte del area de Tecnologia de Información en el siguiente correo adjuntamos el Acta de entrega del equipo que se le entrego a su persona, que detalla lo siguiente:'
            
            from_email = 'soporteic@continental.edu.pe'
            recipient_list = [email]
            email_message = EmailMessage(subject, message, from_email, recipient_list)
        
            if request.FILES:
                for file in request.FILES.getlist('pdf_file'):
                    email_message.attach(file.name, file.read(), file.content_type)

            email_message.send()
        
        return render(request, 'email/email.html')
    return render(request, 'email/email.html')

